"""Отрисовка отчёта в Word, Excel и PDF.

Модуль ничего не считает и не ходит в базу. На вход приходит готовый
:class:`~app.services.reports.ReportDocument`, на выход уходят байты файла. Так
три формата одного отчёта гарантированно содержат одни и те же цифры: если бы
каждый отрисовщик сам доставал значения, расхождение между Word и Excel было бы
вопросом времени, а обнаружилось бы на совещании.

Ячейка приходит уже отформатированной (:class:`~app.services.reports.Cell`).
Отрисовщик не превращает `None` в пустоту и не подставляет ноль — он печатает
`cell.text`. Единственное исключение — Excel: там в числовых колонках пишется
`cell.number`, если оно есть, потому что книгу открывают, чтобы сортировать и
складывать. Когда числа нет, в ячейку уходит тот же текст «нет данных», и
пустой ячейки не возникает ни при каких условиях.

Порядок блоков одинаков во всех форматах и выбран не случайно: предупреждение о
полноте идёт **до** данных. Сноска под таблицей на четвёртой странице не
предупреждение, а алиби.

Про PDF — отдельно, см. :func:`render_pdf`.
"""

from __future__ import annotations

import io
import re
from collections.abc import Iterable, Sequence
from enum import StrEnum
from pathlib import Path
from typing import Any, Final

from app.core.config import get_settings
from app.services.reports import (
    WARNING_HEADING,
    Cell,
    ReportColumn,
    ReportDocument,
    ReportSection,
    ReportTable,
)


class ReportFormat(StrEnum):
    """Формат выгрузки.

    ТЗ (раздел 17) перечисляет ещё PNG и JPEG, но растровые форматы относятся
    к выгрузке **карты**, а не отчёта: картинка не может содержать таблицу
    источников и перечень фильтров так, чтобы их можно было прочитать и
    скопировать. Здесь — три документных формата.
    """

    DOCX = "docx"
    XLSX = "xlsx"
    PDF = "pdf"

    @property
    def media_type(self) -> str:
        return _MEDIA_TYPES[self]

    @property
    def extension(self) -> str:
        return f".{self.value}"


_MEDIA_TYPES: Final[dict[ReportFormat, str]] = {
    ReportFormat.DOCX: (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    ReportFormat.XLSX: "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ReportFormat.PDF: "application/pdf",
}


class PdfUnavailableError(RuntimeError):
    """PDF собрать нечем — и об этом надо сказать, а не отдать битый файл.

    Возбуждается в двух случаях: не установлен `reportlab` или в репозитории
    нет файла шрифта с кириллицей. Оба случая одинаково фатальны для PDF и
    одинаково безобидны для Word и Excel, поэтому исключение отдельное:
    эндпоинт превращает его в 501 с объяснением, а остальные форматы
    продолжают работать.
    """


# --- Шрифт для PDF -----------------------------------------------------------

#: Имена, под которыми шрифты регистрируются в reportlab.
_PDF_FONT: Final = "DejaVuSans"
_PDF_FONT_BOLD: Final = "DejaVuSans-Bold"

#: Файлы шрифта в репозитории. Происхождение и лицензия — data/fonts/PROVENANCE.md.
_FONT_FILES: Final[tuple[tuple[str, str], ...]] = (
    (_PDF_FONT, "DejaVuSans.ttf"),
    (_PDF_FONT_BOLD, "DejaVuSans-Bold.ttf"),
)

#: Признак того, что шрифты уже зарегистрированы в этом процессе. Повторная
#: регистрация не ошибка, но она заново читает 1,4 МБ с диска на каждый отчёт.
_fonts_registered = False


def font_directory() -> Path:
    return Path(get_settings().data_dir) / "fonts"


def register_pdf_fonts() -> None:
    """Зарегистрировать кириллический шрифт в reportlab.

    Шрифт берётся **только** из репозитория. Ни загрузки из сети, ни поиска по
    системным каталогам: ТЗ (раздел 19) требует развёртывания в закрытом
    контуре, где сети нет, а состав системных шрифтов на машине разработчика и
    на сервере заказчика разный — отчёт, который собрался у одного и осыпался
    у другого, хуже отчёта, который не собирается ни у кого.
    """
    global _fonts_registered
    if _fonts_registered:
        return

    try:
        # Типовых заглушек у reportlab нет, поэтому импорт помечен для
        # анализатора. Всё, что берётся из пакета, используется как `Any`, и
        # проверять здесь по существу нечего.
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError as exc:  # pragma: no cover — зависит от окружения
        raise PdfUnavailableError(
            "Для выгрузки в PDF нужен пакет reportlab. Установите его "
            "(pip install reportlab) либо выгрузите отчёт в Word или Excel."
        ) from exc

    directory = font_directory()
    for name, file_name in _FONT_FILES:
        path = directory / file_name
        if not path.is_file():
            raise PdfUnavailableError(
                f"Не найден файл шрифта {file_name}. Ожидался по пути {path}. "
                "Без кириллического шрифта PDF собрался бы нечитаемым, поэтому "
                "выгрузка остановлена. Происхождение и лицензия шрифта описаны "
                "в data/fonts/PROVENANCE.md; Word и Excel работают без него."
            )
        pdfmetrics.registerFont(TTFont(name, str(path)))

    from reportlab.lib.fonts import addMapping

    # Курсив и полужирный курсив отображаются на те же начертания: отдельных
    # файлов мы не возим, а несопоставленное начертание reportlab молча
    # заменит на Helvetica — то есть на шрифт без кириллицы.
    addMapping(_PDF_FONT, 0, 0, _PDF_FONT)
    addMapping(_PDF_FONT, 1, 0, _PDF_FONT_BOLD)
    addMapping(_PDF_FONT, 0, 1, _PDF_FONT)
    addMapping(_PDF_FONT, 1, 1, _PDF_FONT_BOLD)

    _fonts_registered = True


def pdf_available() -> bool:
    """Можно ли собрать PDF в этом развёртывании."""
    try:
        register_pdf_fonts()
    except PdfUnavailableError:
        return False
    return True


# --- Общие блоки -------------------------------------------------------------

_META_TITLE: Final = "Сведения о формировании"
_FILTERS_TITLE: Final = "Применённые фильтры"
_SOURCES_TITLE: Final = "Источники данных"
_NOTES_TITLE: Final = "Примечания"

_SOURCES_EMPTY: Final = (
    "Сведения об источниках в системе отсутствуют: происхождение данных этого "
    "отчёта не подтверждено. Так бывает, если данные загружены в обход "
    "процедуры импорта. Пользуйтесь отчётом с этой оговоркой."
)


def _meta_rows(document: ReportDocument) -> tuple[tuple[str, str], ...]:
    """Обязательная по ТЗ шапка: когда, кем и в какой роли сформировано."""
    return (
        ("Дата и время формирования", document.generated_at_text),
        ("Сформировал", document.generated_by_name),
        ("Роль", document.generated_by_role),
        ("Шаблон отчёта", f"{document.title} ({document.template.value})"),
    )


def _source_rows(document: ReportDocument) -> tuple[tuple[str, ...], ...]:
    return tuple(
        (
            source.layer_code,
            source.file_name,
            ", ".join(source.sheet_names) if source.sheet_names else "—",
            source.as_of_text,
            str(source.row_count) if source.row_count is not None else "нет данных",
        )
        for source in document.sources
    )


_SOURCE_HEADERS: Final = ("Слой", "Файл-источник", "Листы", "Дата актуальности", "Строк")


# --- Word --------------------------------------------------------------------


def render_docx(document: ReportDocument) -> bytes:
    """Собрать отчёт в Word."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading(document.title, level=0)
    subtitle = doc.add_paragraph(document.subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading(_META_TITLE, level=1)
    _docx_pairs(doc, _meta_rows(document))

    # Предупреждение о полноте — до данных и заметно. Читатель, пролиставший
    # отчёт до таблиц, обязан пройти через него.
    doc.add_heading(WARNING_HEADING, level=1)
    for index, line in enumerate(document.warning.lines):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        if index == 0 and document.warning.has_gaps:
            run.bold = True
            run.font.color.rgb = RGBColor(0xB4, 0x1E, 0x1E)

    doc.add_heading(_FILTERS_TITLE, level=1)
    _docx_pairs(doc, document.filters)

    doc.add_heading(_SOURCES_TITLE, level=1)
    if document.sources:
        _docx_table(doc, _SOURCE_HEADERS, _source_rows(document))
    else:
        doc.add_paragraph(_SOURCES_EMPTY)

    for section in document.sections:
        doc.add_heading(section.title, level=1)
        for text in section.paragraphs:
            doc.add_paragraph(text)
        for table in section.tables:
            _docx_report_table(doc, table)

    if document.notes:
        doc.add_heading(_NOTES_TITLE, level=1)
        for note in document.notes:
            paragraph = doc.add_paragraph(note)
            paragraph.runs[0].font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_pairs(doc: Any, pairs: Sequence[tuple[str, str]]) -> None:
    """Таблица «название → значение» в две колонки."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for name, value in pairs:
        cells = table.add_row().cells
        cells[0].text = name
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value


def _docx_table(doc: Any, headers: Sequence[str], rows: Iterable[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def _docx_report_table(doc: Any, table: ReportTable) -> None:
    from docx.shared import Pt

    heading = doc.add_paragraph(table.title)
    heading.runs[0].bold = True

    if table.is_empty:
        # Пустая таблица заменяется фразой, а не пропускается: пропуск читается
        # как «раздел забыли», а не как «строк нет».
        doc.add_paragraph("Строк, удовлетворяющих условиям, нет.")
    else:
        _docx_table(
            doc,
            [column.title for column in table.columns],
            [[cell.text for cell in row] for row in table.rows],
        )

    if table.note:
        paragraph = doc.add_paragraph(table.note)
        paragraph.runs[0].font.size = Pt(8)
        paragraph.runs[0].italic = True


# --- Excel -------------------------------------------------------------------

#: Excel запрещает эти знаки в названии листа и обрезает название до 31 знака.
_SHEET_FORBIDDEN: Final = re.compile(r"[:\\/?*\[\]]")


def render_xlsx(document: ReportDocument) -> bytes:
    """Собрать отчёт в Excel.

    Первый лист — титульный: шапка, предупреждение о полноте, фильтры,
    источники. Дальше по листу на таблицу. Раскладывать всё на один лист нельзя:
    таблицы имеют разное число колонок, и автофильтр в такой книге бесполезен.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = Workbook()
    cover = workbook.active
    assert cover is not None  # у новой книги активный лист есть всегда
    cover.title = "Титул"

    bold = Font(bold=True)
    used: set[str] = {"Титул"}
    row = 1

    def write(values: Sequence[Any], *, header: bool = False) -> None:
        nonlocal row
        for column, value in enumerate(values, start=1):
            target = cover.cell(row=row, column=column, value=value)
            if header:
                target.font = bold
            target.alignment = Alignment(vertical="top", wrap_text=True)
        row += 1

    cover.cell(row=1, column=1, value=document.title).font = Font(bold=True, size=14)
    row = 2
    write((document.subtitle,))
    row += 1

    write((_META_TITLE,), header=True)
    for name, value in _meta_rows(document):
        write((name, value))
    row += 1

    write((WARNING_HEADING,), header=True)
    warning_fill = PatternFill("solid", fgColor="FFF3CD")
    for index, line in enumerate(document.warning.lines):
        cell = cover.cell(row=row, column=1, value=line)
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        if index == 0 and document.warning.has_gaps:
            cell.font = bold
            cell.fill = warning_fill
        row += 1
    row += 1

    write((_FILTERS_TITLE,), header=True)
    for name, value in document.filters:
        write((name, value))
    row += 1

    write((_SOURCES_TITLE,), header=True)
    if document.sources:
        write(_SOURCE_HEADERS, header=True)
        for source_row in _source_rows(document):
            write(source_row)
    else:
        write((_SOURCES_EMPTY,))
    row += 1

    if document.notes:
        write((_NOTES_TITLE,), header=True)
        for note in document.notes:
            write((note,))

    cover.column_dimensions["A"].width = 34
    cover.column_dimensions["B"].width = 80

    for section in document.sections:
        for table in section.tables:
            _xlsx_table_sheet(workbook, section, table, used)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _sheet_name(title: str, used: set[str]) -> str:
    cleaned = _SHEET_FORBIDDEN.sub(" ", title).strip() or "Таблица"
    candidate = cleaned[:31]
    suffix = 2
    while candidate in used:
        tail = f" {suffix}"
        candidate = cleaned[: 31 - len(tail)] + tail
        suffix += 1
    used.add(candidate)
    return candidate


def _xlsx_table_sheet(
    workbook: Any, section: ReportSection, table: ReportTable, used: set[str]
) -> None:
    from openpyxl.styles import Alignment, Font

    sheet = workbook.create_sheet(_sheet_name(table.title, used))
    bold = Font(bold=True)

    sheet.cell(row=1, column=1, value=f"{section.title} — {table.title}").font = bold

    for index, column in enumerate(table.columns, start=1):
        cell = sheet.cell(row=2, column=index, value=column.title)
        cell.font = bold
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        sheet.column_dimensions[cell.column_letter].width = max(12, column.width * 14)

    for row_index, row in enumerate(table.rows, start=3):
        for column_index, (cell, column) in enumerate(
            zip(row, table.columns, strict=False), start=1
        ):
            sheet.cell(row=row_index, column=column_index, value=_xlsx_value(cell, column))
            if column.numeric and cell.number is not None and cell.number_format:
                sheet.cell(row=row_index, column=column_index).number_format = (
                    cell.number_format
                )

    note_row = 3 + len(table.rows) + 1
    if table.is_empty:
        sheet.cell(row=3, column=1, value="Строк, удовлетворяющих условиям, нет.")
        note_row = 5
    if table.note:
        sheet.cell(row=note_row, column=1, value=table.note).alignment = Alignment(
            wrap_text=True
        )


def _xlsx_value(cell: Cell, column: ReportColumn) -> str | float:
    """Что положить в ячейку книги.

    Число — только в числовой колонке и только если оно есть. Во всех
    остальных случаях уходит текст, в том числе «нет данных»: пустая ячейка в
    книге неотличима от нуля, и именно на этом спотыкались исходные книги
    заказчика.
    """
    if column.numeric and cell.number is not None:
        return cell.number
    return cell.text


# --- PDF ---------------------------------------------------------------------


def render_pdf(document: ReportDocument) -> bytes:
    """Собрать отчёт в PDF.

    Решение по стеку. WeasyPrint отвергнут: на Windows он требует GTK,
    поставляемого отдельно от Python, а система разрабатывается и
    разворачивается в том числе на Windows. `fpdf2` в зависимостях нет и он
    заметно беднее по вёрстке таблиц. Взят `reportlab` — он умеет
    встраивать TrueType, а вёрстка таблиц у него встроенная.

    Кириллица. Все четырнадцать встроенных шрифтов PDF кодируются в WinAnsi и
    кириллицы не содержат; шрифты из комплекта reportlab (семейство Vera) её
    тоже не содержат — проверено по таблице cmap. Поэтому в репозитории лежит
    DejaVu Sans под свободной лицензией (data/fonts/PROVENANCE.md), и он
    регистрируется из файла. Если файла нет — :class:`PdfUnavailableError`, а не
    молчаливый откат на Helvetica: откат дал бы файл, который открывается,
    выглядит сформированным и не читается.
    """
    register_pdf_fonts()

    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import (
        ParagraphStyle,
        getSampleStyleSheet,
    )
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    base = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "RiskTitle", parent=base["Title"], fontName=_PDF_FONT_BOLD, fontSize=16, leading=20
    )
    heading_style = ParagraphStyle(
        "RiskHeading",
        parent=base["Heading2"],
        fontName=_PDF_FONT_BOLD,
        fontSize=11,
        leading=14,
        spaceBefore=8,
        alignment=TA_LEFT,
    )
    body_style = ParagraphStyle(
        "RiskBody", parent=base["BodyText"], fontName=_PDF_FONT, fontSize=8.5, leading=11
    )
    alarm_style = ParagraphStyle(
        "RiskAlarm",
        parent=body_style,
        fontName=_PDF_FONT_BOLD,
        textColor=colors.HexColor("#B41E1E"),
    )
    cell_style = ParagraphStyle(
        "RiskCell", parent=body_style, fontSize=7, leading=8.5, spaceAfter=0
    )
    header_cell_style = ParagraphStyle(
        "RiskCellHead", parent=cell_style, fontName=_PDF_FONT_BOLD
    )
    note_style = ParagraphStyle(
        "RiskNote", parent=body_style, fontSize=7, leading=9, textColor=colors.grey
    )

    buffer = io.BytesIO()
    page = landscape(A4)
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=page,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=document.title,
        author=document.generated_by_name,
    )
    available = page[0] - pdf.leftMargin - pdf.rightMargin

    story: list[Any] = [
        Paragraph(_escape(document.title), title_style),
        Paragraph(_escape(document.subtitle), body_style),
        Spacer(1, 6),
        Paragraph(_escape(_META_TITLE), heading_style),
        _pdf_pairs(_meta_rows(document), available, cell_style, header_cell_style, colors),
        Paragraph(_escape(WARNING_HEADING), heading_style),
    ]

    for index, line in enumerate(document.warning.lines):
        style = alarm_style if index == 0 and document.warning.has_gaps else body_style
        story.append(Paragraph(_escape(line), style))

    story.append(Paragraph(_escape(_FILTERS_TITLE), heading_style))
    story.append(
        _pdf_pairs(document.filters, available, cell_style, header_cell_style, colors)
    )

    story.append(Paragraph(_escape(_SOURCES_TITLE), heading_style))
    if document.sources:
        story.append(
            _pdf_grid(
                [list(_SOURCE_HEADERS)] + [list(row) for row in _source_rows(document)],
                [0.6, 3.0, 2.0, 1.4, 0.8],
                available,
                cell_style,
                header_cell_style,
                colors,
            )
        )
    else:
        story.append(Paragraph(_escape(_SOURCES_EMPTY), body_style))

    for section in document.sections:
        story.append(PageBreak())
        story.append(Paragraph(_escape(section.title), heading_style))
        for paragraph in section.paragraphs:
            story.append(Paragraph(_escape(paragraph), body_style))
        for table in section.tables:
            story.append(Spacer(1, 4))
            story.append(Paragraph(_escape(table.title), heading_style))
            if table.is_empty:
                story.append(
                    Paragraph("Строк, удовлетворяющих условиям, нет.", body_style)
                )
            else:
                story.append(
                    _pdf_grid(
                        [[column.title for column in table.columns]]
                        + [[cell.text for cell in row] for row in table.rows],
                        [column.width for column in table.columns],
                        available,
                        cell_style,
                        header_cell_style,
                        colors,
                    )
                )
            if table.note:
                story.append(Paragraph(_escape(table.note), note_style))

    if document.notes:
        story.append(Spacer(1, 8))
        block: list[Any] = [Paragraph(_escape(_NOTES_TITLE), heading_style)]
        block.extend(Paragraph(_escape(note), note_style) for note in document.notes)
        story.append(KeepTogether(block))

    pdf.build(story, onFirstPage=_pdf_footer(document), onLaterPages=_pdf_footer(document))
    return buffer.getvalue()


def _pdf_footer(document: ReportDocument) -> Any:
    """Колонтитул: кто и когда сформировал, плюс номер страницы.

    Отчёт распечатывают и разносят по кабинетам постранично. Страница без
    указания источника и даты через неделю становится листком неизвестного
    происхождения.
    """

    def draw(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont(_PDF_FONT, 7)
        canvas.setFillGray(0.4)
        left = (
            f"{document.title} · сформировал: {document.generated_by_name} "
            f"({document.generated_by_role}) · {document.generated_at_text}"
        )
        canvas.drawString(doc.leftMargin, 8 * 72 / 25.4, left)
        canvas.drawRightString(
            doc.pagesize[0] - doc.rightMargin, 8 * 72 / 25.4, f"стр. {doc.page}"
        )
        canvas.restoreState()

    return draw


def _escape(text: str) -> str:
    """Экранирование для мини-разметки reportlab.

    Названия организаций из источников содержат `&` и угловые скобки, и без
    экранирования reportlab примет их за разметку и уронит сборку документа —
    на отдельно взятой строке из четырёх тысяч.
    """
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_pairs(
    pairs: Sequence[tuple[str, str]],
    available: float,
    cell_style: Any,
    header_style: Any,
    colors: Any,
) -> Any:
    return _pdf_grid(
        [[name, value] for name, value in pairs],
        [1.2, 4.0],
        available,
        cell_style,
        header_style,
        colors,
        with_header=False,
    )


def _pdf_grid(
    rows: Sequence[Sequence[str]],
    widths: Sequence[float],
    available: float,
    cell_style: Any,
    header_style: Any,
    colors: Any,
    *,
    with_header: bool = True,
) -> Any:
    from reportlab.platypus import Paragraph, Table, TableStyle

    total = sum(widths) or 1.0
    column_widths = [available * width / total for width in widths]

    data = [
        [
            Paragraph(
                _escape(str(value)),
                header_style if with_header and index == 0 else cell_style,
            )
            for value in row
        ]
        for index, row in enumerate(rows)
    ]

    table = Table(data, colWidths=column_widths, repeatRows=1 if with_header else 0)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B9C0C8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if with_header:
        style.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDF1F5")))
    else:
        style.append(("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5F7FA")))
    table.setStyle(TableStyle(style))
    return table


# --- Единая точка входа ------------------------------------------------------

_RENDERERS: Final[dict[ReportFormat, Any]] = {
    ReportFormat.DOCX: render_docx,
    ReportFormat.XLSX: render_xlsx,
    ReportFormat.PDF: render_pdf,
}


def render(document: ReportDocument, report_format: ReportFormat) -> bytes:
    """Отрисовать отчёт в заданном формате."""
    renderer = _RENDERERS[report_format]
    result: bytes = renderer(document)
    return result


__all__ = [
    "PdfUnavailableError",
    "ReportFormat",
    "font_directory",
    "pdf_available",
    "register_pdf_fonts",
    "render",
    "render_docx",
    "render_pdf",
    "render_xlsx",
]
